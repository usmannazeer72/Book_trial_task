"""
Compilation Service - Compile final book from chapters
"""

import os
import logging
from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from src.database.schema import DatabaseManager
from src.services.notification_service import NotificationService
from src.config.settings import config

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


class CompilationService:
    """Service to compile chapters into final document"""
    
    def __init__(self):
        self.db_manager = DatabaseManager()
        self.db = self.db_manager.connect()
        self.notification_service = NotificationService()
        self.output_dir = config.app.output_directory
        
        # Ensure output directory exists
        os.makedirs(self.output_dir, exist_ok=True)
    
    def can_compile(self, book_id: str) -> tuple[bool, str]:
        """
        Check if book can be compiled based on gating logic
        
        Returns:
            (can_compile, reason)
        """
        book = self.db.books.find_one({'book_id': book_id})
        
        if not book:
            return False, "Book not found"
        
        # Check final review status
        final_status = book.get('final_review_notes_status', '')
        
        if final_status == 'no':
            return False, "Final compilation paused by editor"
        
        if final_status == 'yes':
            # Editor wants to add notes - check if notes exist
            # For now, we'll allow compilation if status is 'yes'
            # In a real system, you might wait for specific final notes
            pass
        
        # Check if outline is approved
        outline_status = book.get('status_outline_notes', '')
        if outline_status != 'no_notes_needed':
            return False, f"Outline not approved (status: {outline_status})"
        
        # Check if all chapters are generated and approved
        chapters = list(self.db.chapters.find({'book_id': book_id}).sort('chapter_number', 1))
        
        if not chapters:
            return False, "No chapters found"
        
        for chapter in chapters:
            chapter_status = chapter.get('chapter_notes_status', '')
            if chapter_status not in ['no_notes_needed', 'pending']:
                return False, f"Chapter {chapter['chapter_number']} not approved (status: {chapter_status})"
        
        # Check if already compiled
        book_status = book.get('book_output_status', '')
        if book_status == 'completed':
            return False, "Book already compiled"
        
        return True, "Ready to compile"
    
    def compile_to_docx(self, book_id: str) -> str:
        """
        Compile book to DOCX format
        
        Returns:
            Path to generated file
        """
        book = self.db.books.find_one({'book_id': book_id})
        outline = self.db.outlines.find_one(
            {'book_id': book_id},
            sort=[('version', -1)]
        )
        chapters = list(self.db.chapters.find({'book_id': book_id}).sort('chapter_number', 1))
        
        title = book['title']
        
        # Create document
        doc = Document()
        
        # Add title page
        title_para = doc.add_paragraph()
        title_run = title_para.add_run(title)
        title_run.font.size = Pt(24)
        title_run.font.bold = True
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add some space
        doc.add_paragraph()
        doc.add_paragraph()
        
        # Add generation date
        date_para = doc.add_paragraph()
        date_run = date_para.add_run(f"Generated: {datetime.utcnow().strftime('%B %d, %Y')}")
        date_run.font.size = Pt(12)
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Page break
        doc.add_page_break()
        
        # Add outline section
        outline_heading = doc.add_heading('Book Outline', level=1)
        doc.add_paragraph(outline['outline_text'])
        doc.add_page_break()
        
        # Add chapters
        for chapter in chapters:
            # Chapter heading
            chapter_heading = doc.add_heading(
                f"Chapter {chapter['chapter_number']}: {chapter['chapter_title']}", 
                level=1
            )
            
            # Chapter content
            content_paragraphs = chapter['content'].split('\n\n')
            for para_text in content_paragraphs:
                if para_text.strip():
                    doc.add_paragraph(para_text.strip())
            
            # Page break after each chapter
            doc.add_page_break()
        
        # Save document
        filename = f"{title.replace(' ', '_')}_{book_id[:8]}.docx"
        filepath = os.path.join(self.output_dir, filename)
        doc.save(filepath)
        
        logger.info(f"✓ DOCX compiled: {filepath}")
        return filepath
    
    def compile_to_pdf(self, book_id: str) -> str:
        """
        Compile book to PDF format
        
        Returns:
            Path to generated file
        """
        book = self.db.books.find_one({'book_id': book_id})
        outline = self.db.outlines.find_one(
            {'book_id': book_id},
            sort=[('version', -1)]
        )
        chapters = list(self.db.chapters.find({'book_id': book_id}).sort('chapter_number', 1))
        
        title = book['title']
        
        # Setup PDF
        filename = f"{title.replace(' ', '_')}_{book_id[:8]}.pdf"
        filepath = os.path.join(self.output_dir, filename)
        
        doc = SimpleDocTemplate(filepath, pagesize=letter)
        story = []
        styles = getSampleStyleSheet()
        
        # Custom styles
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Title'],
            fontSize=24,
            spaceAfter=30,
            alignment=1  # Center
        )
        
        heading_style = ParagraphStyle(
            'CustomHeading',
            parent=styles['Heading1'],
            fontSize=18,
            spaceAfter=12,
            spaceBefore=12
        )
        
        # Title page
        story.append(Paragraph(title, title_style))
        story.append(Spacer(1, 0.5*inch))
        story.append(Paragraph(f"Generated: {datetime.utcnow().strftime('%B %d, %Y')}", styles['Normal']))
        story.append(PageBreak())
        
        # Outline
        story.append(Paragraph("Book Outline", heading_style))
        outline_paragraphs = outline['outline_text'].split('\n\n')
        for para in outline_paragraphs:
            if para.strip():
                story.append(Paragraph(para.strip(), styles['Normal']))
                story.append(Spacer(1, 0.2*inch))
        story.append(PageBreak())
        
        # Chapters
        for chapter in chapters:
            # Chapter heading
            chapter_title = f"Chapter {chapter['chapter_number']}: {chapter['chapter_title']}"
            story.append(Paragraph(chapter_title, heading_style))
            story.append(Spacer(1, 0.2*inch))
            
            # Chapter content
            content_paragraphs = chapter['content'].split('\n\n')
            for para in content_paragraphs:
                if para.strip():
                    story.append(Paragraph(para.strip(), styles['Normal']))
                    story.append(Spacer(1, 0.1*inch))
            
            story.append(PageBreak())
        
        # Build PDF
        doc.build(story)
        
        logger.info(f"✓ PDF compiled: {filepath}")
        return filepath
    
    def compile_to_txt(self, book_id: str) -> str:
        """
        Compile book to plain text format
        
        Returns:
            Path to generated file
        """
        book = self.db.books.find_one({'book_id': book_id})
        outline = self.db.outlines.find_one(
            {'book_id': book_id},
            sort=[('version', -1)]
        )
        chapters = list(self.db.chapters.find({'book_id': book_id}).sort('chapter_number', 1))
        
        title = book['title']
        
        # Build text content
        content = []
        content.append("=" * 80)
        content.append(title.upper().center(80))
        content.append("=" * 80)
        content.append("")
        content.append(f"Generated: {datetime.utcnow().strftime('%B %d, %Y')}")
        content.append("")
        content.append("=" * 80)
        content.append("")
        
        # Outline
        content.append("BOOK OUTLINE")
        content.append("-" * 80)
        content.append(outline['outline_text'])
        content.append("")
        content.append("=" * 80)
        content.append("")
        
        # Chapters
        for chapter in chapters:
            content.append(f"CHAPTER {chapter['chapter_number']}: {chapter['chapter_title'].upper()}")
            content.append("-" * 80)
            content.append("")
            content.append(chapter['content'])
            content.append("")
            content.append("=" * 80)
            content.append("")
        
        # Save to file
        filename = f"{title.replace(' ', '_')}_{book_id[:8]}.txt"
        filepath = os.path.join(self.output_dir, filename)
        
        with open(filepath, 'w', encoding='utf-8') as f:
            f.write('\n'.join(content))
        
        logger.info(f"✓ TXT compiled: {filepath}")
        return filepath
    
    def compile_book(self, book_id: str, output_formats: list = None) -> dict:
        """
        Compile book in specified formats
        
        Args:
            book_id: Book ID
            output_formats: List of formats ['docx', 'pdf', 'txt']
                           Defaults to all formats
        
        Returns:
            Dict with format: filepath mappings
        """
        if output_formats is None:
            output_formats = ['docx', 'pdf', 'txt']
        
        try:
            # Check if we can compile
            can_compile, reason = self.can_compile(book_id)
            
            if not can_compile:
                logger.warning(f"Cannot compile book {book_id}: {reason}")
                return {}
            
            book = self.db.books.find_one({'book_id': book_id})
            title = book['title']
            
            logger.info(f"Compiling book: {title}")
            
            # Update status to in_progress
            self.db.books.update_one(
                {'book_id': book_id},
                {
                    '$set': {
                        'book_output_status': 'in_progress',
                        'updated_at': datetime.utcnow()
                    }
                }
            )
            
            output_files = {}
            
            # Compile in each requested format
            if 'docx' in output_formats:
                output_files['docx'] = self.compile_to_docx(book_id)
            
            if 'pdf' in output_formats:
                output_files['pdf'] = self.compile_to_pdf(book_id)
            
            if 'txt' in output_formats:
                output_files['txt'] = self.compile_to_txt(book_id)
            
            # Update book status to completed
            self.db.books.update_one(
                {'book_id': book_id},
                {
                    '$set': {
                        'book_output_status': 'completed',
                        'updated_at': datetime.utcnow()
                    }
                }
            )
            
            logger.info(f"✓ Book compilation complete: {title}")
            
            # Send notification
            self.notification_service.notify_book_completed(
                book_id,
                title,
                ', '.join(output_files.values())
            )
            
            return output_files
        
        except Exception as e:
            logger.error(f"Error compiling book {book_id}: {str(e)}")
            
            # Update status to failed
            self.db.books.update_one(
                {'book_id': book_id},
                {
                    '$set': {
                        'book_output_status': 'failed',
                        'updated_at': datetime.utcnow()
                    }
                }
            )
            
            # Notify about error
            book = self.db.books.find_one({'book_id': book_id})
            if book:
                self.notification_service.notify_error(
                    book_id,
                    book['title'],
                    f"Error compiling book: {str(e)}"
                )
            
            return {}
    
    def close(self):
        """Close all connections"""
        self.db_manager.close()
        self.notification_service.close()


if __name__ == "__main__":
    # Test compilation service
    service = CompilationService()
    
    # Example: Compile a specific book
    book_id = "your_book_id_here"
    output_files = service.compile_book(book_id)
    print(f"Compiled files: {output_files}")
    
    service.close()
